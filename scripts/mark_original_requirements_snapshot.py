#!/usr/bin/env python3
"""Mark the original customer requirements DOCX as the F-57 source baseline.

The original wording is preserved verbatim below the inserted status notice so
that later scope decisions remain auditable. The script is intentionally
idempotent.
"""

import os
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from docx_package_hygiene import sanitize_docx_package


WORKSPACE = Path(__file__).resolve().parents[1]
DEFAULT_INPUT = WORKSPACE / "docs" / "介绍" / "管理软件基本需求.docx"
INPUT = Path(os.environ.get("REQUIREMENTS_SNAPSHOT_INPUT", DEFAULT_INPUT)).expanduser()
OUTPUT = Path(os.environ.get("REQUIREMENTS_SNAPSHOT_OUTPUT", INPUT)).expanduser()

# 冻结时间戳：docx 的 core properties 与 ZIP 条目时间都取它，两次运行才逐字节相同。
DOCUMENT_TIMESTAMP = datetime(2026, 8, 26, 0, 0, 0, tzinfo=timezone.utc)
PACKAGE_TIMESTAMP = (
    DOCUMENT_TIMESTAMP.year,
    DOCUMENT_TIMESTAMP.month,
    DOCUMENT_TIMESTAMP.day,
    DOCUMENT_TIMESTAMP.hour,
    DOCUMENT_TIMESTAMP.minute,
    DOCUMENT_TIMESTAMP.second,
)
OLD_MARKER = "文档状态：原始需求输入（历史快照）"
MARKER = "文档状态：客户原始业务需求基线（F-57 追踪来源）"
NOTE = (
    "本文件原样保留客户业务需求，是 F-57 需求追踪的正式来源之一；下方原文不直接定义技术架构、首发范围或安全实现。"
    "现行范围、接口、数据、权限、安全、测试与部署口径，以仓库 README 指向的 F-57 总体设计、需求追踪、权威登记和实施计划为准。"
)
SCOPE = (
    "解释规则：下方“销售、采购、财务、技术、管理者”是默认人物与工作台模板，不是写死岗位；权限由主体、能力、数据范围、条件、期限、设备、金额和状态动态决定。"
    "“生产触发采购”作为外部生产需求接口保留，完整 MRP/MES 延期；本地模型实现延期，MCP 与服务器控制中心按 F-57 的受控能力和权威边界实现。"
)
LATIN_FONT = "Arial"
CHINESE_FONT = "Microsoft YaHei"
WESTERN_LANG = "en-US"
EAST_ASIA_LANG = "zh-CN"
BIDI_LANG = "ar-SA"
SECTION_HEADINGS = {
    "销售",
    "采购：",
    "财务：",
    "工单系统",
    "管理",
    "合同管理",
    "订单管理",
    "采购与供应商",
    "报表",
    "高度定制，但不破坏升级",
}


def set_language_properties(rpr) -> None:
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), WESTERN_LANG)
    lang.set(qn("w:eastAsia"), EAST_ASIA_LANG)
    lang.set(qn("w:bidi"), BIDI_LANG)


def set_font_properties(fonts, *, east_asia_font: str = CHINESE_FONT) -> None:
    for attribute in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        fonts.attrib.pop(qn(f"w:{attribute}"), None)
    fonts.set(qn("w:ascii"), LATIN_FONT)
    fonts.set(qn("w:hAnsi"), LATIN_FONT)
    fonts.set(qn("w:eastAsia"), east_asia_font)
    fonts.set(qn("w:cs"), LATIN_FONT)


def set_cjk_font(run, name: str = CHINESE_FONT) -> None:
    run.font.name = LATIN_FONT
    rpr = run._element.get_or_add_rPr()
    set_font_properties(rpr.rFonts, east_asia_font=name)
    set_language_properties(rpr)


def configure_document_defaults(doc: Document) -> None:
    styles = doc.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)
    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        doc_defaults.insert(0, rpr_default)
    rpr = rpr_default.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_default.append(rpr)
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    set_font_properties(fonts)
    set_language_properties(rpr)


def iter_paragraphs(doc: Document):
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def finalize_document(doc: Document) -> None:
    configure_document_defaults(doc)
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet"):
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = LATIN_FONT
        rpr = style._element.get_or_add_rPr()
        set_font_properties(rpr.rFonts)
        set_language_properties(rpr)

    for paragraph in iter_paragraphs(doc):
        text = paragraph.text.strip()
        if text in SECTION_HEADINGS:
            paragraph.style = doc.styles["Heading 1"]
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
        for run in paragraph.runs:
            set_cjk_font(run)

    props = doc.core_properties
    props.title = "管理软件基本需求（客户原始业务需求基线）"
    props.subject = "客户原始业务需求基线（F-57 追踪来源）"
    props.keywords = "客户需求, 企业管理, CRM, 合同, 订单, 采购, 财务, 售后, F-57"
    props.author = "客户原始需求（F-57 基线）"
    props.last_modified_by = "企业一体化经营管理平台"
    # 冻结时间戳：`OUTPUT` 默认就是 `INPUT`（本文件 :25），用 now() 会让每次运行都改写
    # 一个受版本控制的 docx，产生永久性的伪差异。兄弟脚本 build_product_intro.py 已用
    # 同一做法做到逐字节可复现（F-68）。
    props.created = DOCUMENT_TIMESTAMP
    props.modified = DOCUMENT_TIMESTAMP


def style_marker_heading(run, paragraph) -> None:
    """标记段标题的样式只此一处。

    原先插入路径给琥珀色 9C6500 加底纹、更新路径给蓝色 1F4E78 不加底纹，
    同一份文档跑两次会得到两种样子——与本模块文档字符串自称的幂等相抵（F-68）。
    """
    run.font.color.rgb = RGBColor(0x9C, 0x65, 0x00)
    shade_paragraph(paragraph)


def shade_paragraph(paragraph, fill: str = "FFF4CE") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    p_pr.append(shd)


def replace_paragraph_text(paragraph, text: str, *, bold: bool = False, size: float = 10.5) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    run = paragraph.add_run(text)
    set_cjk_font(run)
    run.font.size = Pt(size)
    run.bold = bold


def build() -> None:
    doc = Document(INPUT)
    marker_index = next(
        (
            index
            for index, paragraph in enumerate(doc.paragraphs[:8])
            if paragraph.text.strip() in {OLD_MARKER, MARKER}
        ),
        None,
    )
    if marker_index is not None:
        heading = doc.paragraphs[marker_index]
        note = doc.paragraphs[marker_index + 1]
        scope = doc.paragraphs[marker_index + 2]
        replace_paragraph_text(heading, MARKER, bold=True, size=14)
        style_marker_heading(heading.runs[0], heading)
        replace_paragraph_text(note, NOTE)
        replace_paragraph_text(scope, SCOPE, bold=True)
        props = doc.core_properties
        props.comments = "原始正文保持不变；顶部状态说明由 F-57 权威收口更新。"
        finalize_document(doc)
        OUTPUT.parent.mkdir(parents=True, exist_ok=True)
        temp_output = OUTPUT.with_suffix(".tmp.docx")
        doc.save(temp_output)
        sanitize_docx_package(temp_output, package_timestamp=PACKAGE_TIMESTAMP)
        os.replace(temp_output, OUTPUT)
        print(OUTPUT)
        return

    first = doc.paragraphs[0]

    heading = first.insert_paragraph_before()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.paragraph_format.space_after = Pt(4)
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run(MARKER)
    set_cjk_font(run)
    run.bold = True
    run.font.size = Pt(14)
    style_marker_heading(run, heading)

    note = first.insert_paragraph_before()
    note.paragraph_format.space_after = Pt(4)
    note.paragraph_format.line_spacing = 1.2
    run = note.add_run(
        NOTE
    )
    set_cjk_font(run)
    run.font.size = Pt(10.5)

    scope = first.insert_paragraph_before()
    scope.paragraph_format.space_after = Pt(10)
    scope.paragraph_format.line_spacing = 1.2
    run = scope.add_run(
        SCOPE
    )
    set_cjk_font(run)
    run.font.size = Pt(10.5)
    run.bold = True

    props = doc.core_properties
    props.comments = "原始正文保持不变；顶部状态说明由 F-57 权威收口加入。"
    finalize_document(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    temp_output = OUTPUT.with_suffix(".tmp.docx")
    doc.save(temp_output)
    sanitize_docx_package(temp_output, package_timestamp=PACKAGE_TIMESTAMP)
    os.replace(temp_output, OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
