#!/usr/bin/env python3
"""Build the customer-facing product introduction and feature outline.

依赖与重建方式：
- 唯一外部依赖是 python-docx，版本范围 >=1.2,<2。
- 建仓库根目录虚拟环境：python3 -m venv .venv-docx
- 安装依赖：.venv-docx/bin/pip install "python-docx>=1.2,<2"
- 重建文档：.venv-docx/bin/python scripts/build_product_intro.py
- 输出默认写入仓库内的 docs/介绍/ 目录，可用环境变量 PRODUCT_INTRO_OUTPUT 覆盖。
"""

import os
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from docx_package_hygiene import sanitize_docx_package


WORKSPACE = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = WORKSPACE / "docs" / "介绍" / "企业一体化经营管理平台-产品介绍与功能大纲.docx"
OUTPUT = Path(os.environ.get("PRODUCT_INTRO_OUTPUT", DEFAULT_OUTPUT)).expanduser()
# Freeze OOXML core and ZIP metadata so identical sources rebuild byte-for-byte.
# Advance this only when the published document baseline is intentionally revised.
DOCUMENT_TIMESTAMP = datetime(2026, 8, 26, 0, 0, 0, tzinfo=timezone.utc)


# launch_messaging_guide -> compact_reference_guide token map.
PAGE_WIDTH_DXA = 9360
# 表格左缩进取 0，使表宽与正文可用宽度一致，右边缘不越过正文右边界。
TABLE_INDENT_DXA = 0
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
LIST_MARKER_DXA = 270
LIST_TEXT_DXA = 540
LIST_HANGING_DXA = 270

# Windows Server 2022 is the production authority host. Use its standard Latin
# face and the Simplified Chinese UI face, while declaring zh-CN on every text
# layer so Word/LibreOffice can choose an installed CJK fallback on other hosts.
LATIN_FONT = "Arial"
CHINESE_FONT = "Microsoft YaHei"
WESTERN_LANG = "en-US"
EAST_ASIA_LANG = "zh-CN"
BIDI_LANG = "ar-SA"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
BODY = "202830"
MUTED = "667085"
LIGHT_FILL = "F4F6F9"
TABLE_FILL = "E8EEF5"
WHITE = "FFFFFF"
BORDER = "C9D5E3"


def _insert_ordered(parent, child, order):
    """Insert child into parent at the position required by the OOXML sequence."""
    name = child.tag.split("}")[-1]
    idx = order.index(name)
    for existing in parent:
        existing_name = existing.tag.split("}")[-1]
        if existing_name in order and order.index(existing_name) > idx:
            existing.addprevious(child)
            return
    parent.append(child)


TC_PR_ORDER = (
    "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders", "shd",
    "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark",
)
TBL_PR_ORDER = (
    "tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
    "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd", "tblBorders",
    "shd", "tblLayout", "tblCellMar", "tblLook",
)
TR_PR_ORDER = (
    "cnfStyle", "divId", "gridBefore", "gridAfter", "wBefore", "wAfter",
    "cantSplit", "trHeight", "tblHeader",
)
P_PR_ORDER = (
    "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr", "widowControl",
    "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs", "suppressAutoHyphens",
    "kinsoku", "wordWrap", "overflowPunct", "topLinePunct", "autoSpaceDE",
    "autoSpaceDN", "bidi", "adjustRightInd", "snapToGrid", "spacing", "ind",
    "contextualSpacing", "mirrorIndents", "suppressOverlap", "jc", "textDirection",
    "textAlignment", "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr",
)


def apply_table_geometry(table, widths, *, table_width_dxa, indent_dxa, cell_margins_dxa):
    """Apply deterministic table geometry without relying on external helpers."""
    if sum(widths) != table_width_dxa:
        raise ValueError(
            f"列宽合计 {sum(widths)} 与表宽 {table_width_dxa} 不一致"
        )
    if table_width_dxa + indent_dxa > PAGE_WIDTH_DXA:
        raise ValueError(
            f"表宽 {table_width_dxa} 与左缩进 {indent_dxa} 之和超出正文可用宽度 {PAGE_WIDTH_DXA}"
        )
    table.autofit = False
    table_pr = table._tbl.tblPr

    tbl_width = table_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        table_pr.insert(0, tbl_width)
    tbl_width.set(qn("w:w"), str(table_width_dxa))
    tbl_width.set(qn("w:type"), "dxa")

    tbl_layout = table_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        _insert_ordered(table_pr, tbl_layout, TBL_PR_ORDER)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_indent = table_pr.find(qn("w:tblInd"))
    if tbl_indent is None:
        tbl_indent = OxmlElement("w:tblInd")
        _insert_ordered(table_pr, tbl_indent, TBL_PR_ORDER)
    tbl_indent.set(qn("w:w"), str(indent_dxa))
    tbl_indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.insert(0, tc_width)
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")

            margins = tc_pr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                _insert_ordered(tc_pr, margins, TC_PR_ORDER)
            for side, value in cell_margins_dxa.items():
                margin = margins.find(qn(f"w:{side}"))
                if margin is None:
                    margin = OxmlElement(f"w:{side}")
                    margins.append(margin)
                margin.set(qn("w:w"), str(value))
                margin.set(qn("w:type"), "dxa")


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_language_properties(rpr) -> None:
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), WESTERN_LANG)
    lang.set(qn("w:eastAsia"), EAST_ASIA_LANG)
    lang.set(qn("w:bidi"), BIDI_LANG)


def set_font_properties(fonts, *, latin_font: str, east_asia_font: str) -> None:
    for attribute in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        fonts.attrib.pop(qn(f"w:{attribute}"), None)
    fonts.set(qn("w:ascii"), latin_font)
    fonts.set(qn("w:hAnsi"), latin_font)
    fonts.set(qn("w:eastAsia"), east_asia_font)
    fonts.set(qn("w:cs"), latin_font)


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: str = BODY,
    bold: bool | None = None,
    italic: bool | None = None,
    latin_font: str = LATIN_FONT,
    east_asia_font: str = CHINESE_FONT,
):
    run.font.name = latin_font
    rpr = run._element.get_or_add_rPr()
    set_font_properties(
        rpr.rFonts,
        latin_font=latin_font,
        east_asia_font=east_asia_font,
    )
    set_language_properties(rpr)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_style_font(style, *, size: float, color: str, bold: bool = False):
    style.font.name = LATIN_FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = rgb(color)
    rpr = style._element.get_or_add_rPr()
    set_font_properties(
        rpr.rFonts,
        latin_font=LATIN_FONT,
        east_asia_font=CHINESE_FONT,
    )
    set_language_properties(rpr)


def configure_document_defaults(doc: Document) -> None:
    """Declare fonts and languages even for runs without direct formatting."""
    styles = doc.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)
    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        doc_defaults.insert(0, rpr_default)
    rpr = rpr_default.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_default.append(rpr)
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    set_font_properties(
        fonts,
        latin_font=LATIN_FONT,
        east_asia_font=CHINESE_FONT,
    )
    set_language_properties(rpr)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, color=BODY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[style_name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name, size, color, bold, before, after, line in (
        ("Intro Lead", 13.5, NAVY, False, 0, 12, 1.25),
        ("Small Note", 9.5, MUTED, False, 0, 4, 1.15),
        ("Callout", 11.5, NAVY, True, 0, 0, 1.25),
    ):
        if name not in doc.styles:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = doc.styles[name]
        set_style_font(style, size=size, color=color, bold=bold)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = line


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    left = p.add_run("企业一体化经营管理平台")
    set_run_font(left, size=9, color=MUTED, bold=True)
    p.add_run("\t")
    right = p.add_run("产品介绍与功能大纲")
    set_run_font(right, size=9, color=MUTED)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""

    for footer in (section.footer, section.first_page_footer):
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        prefix = p.add_run("第 ")
        set_run_font(prefix, size=9, color=MUTED)
        add_page_field(p)
        suffix = p.add_run(" 页")
        set_run_font(suffix, size=9, color=MUTED)


def next_numbering_id(numbering_root, tag_name: str) -> int:
    values = []
    for element in numbering_root.findall(qn(tag_name)):
        attr = "w:abstractNumId" if tag_name == "w:abstractNum" else "w:numId"
        raw = element.get(qn(attr))
        if raw is not None:
            values.append(int(raw))
    return max(values, default=0) + 1


def create_numbering(doc: Document, *, decimal: bool) -> int:
    root = doc.part.numbering_part.element
    abstract_id = next_numbering_id(root, "w:abstractNum")
    num_id = next_numbering_id(root, "w:num")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if decimal else "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if decimal else "•")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(LIST_TEXT_DXA))
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(LIST_TEXT_DXA))
    ind.set(qn("w:hanging"), str(LIST_HANGING_DXA))
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    ppr.append(ind)
    lvl.append(ppr)

    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    set_font_properties(
        fonts,
        latin_font=LATIN_FONT,
        east_asia_font=CHINESE_FONT,
    )
    rpr.append(fonts)
    set_language_properties(rpr)
    lvl.append(rpr)
    abstract.append(lvl)
    # OOXML CT_Numbering 要求全部 w:abstractNum 排在 w:num 之前，追加到末尾会破坏部件结构
    first_num = root.find(qn("w:num"))
    if first_num is None:
        root.append(abstract)
    else:
        first_num.addprevious(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    root.append(num)
    return num_id


def apply_num(paragraph, num_id: int, *, compact: bool = False):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_after = Pt(2 if compact else 4)
    paragraph.paragraph_format.line_spacing = 1.15 if compact else 1.25


def add_bullet(
    doc: Document,
    text: str,
    bullet_num_id: int,
    *,
    bold_lead: str | None = None,
    compact: bool = False,
    font_size: float | None = None,
    line_spacing: float | None = None,
):
    p = doc.add_paragraph()
    apply_num(p, bullet_num_id, compact=compact)
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing
    font_size = font_size if font_size is not None else (10.5 if compact else 11)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=font_size, color=BODY, bold=True)
        rest = p.add_run(text[len(bold_lead) :])
        set_run_font(rest, size=font_size, color=BODY)
    else:
        run = p.add_run(text)
        set_run_font(run, size=font_size, color=BODY)
    return p


def add_numbered(
    doc: Document,
    text: str,
    number_num_id: int,
    *,
    bold_lead: str | None = None,
    compact: bool = False,
):
    p = doc.add_paragraph()
    apply_num(p, number_num_id, compact=compact)
    font_size = 10.5 if compact else 11
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=font_size, color=NAVY, bold=True)
        rest = p.add_run(text[len(bold_lead) :])
        set_run_font(rest, size=font_size, color=BODY)
    else:
        run = p.add_run(text)
        set_run_font(run, size=font_size, color=BODY)
    return p


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        _insert_ordered(tc_pr, shd, TC_PR_ORDER)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_borders(cell, color: str = BORDER, size: int = 6):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        _insert_ordered(tc_pr, borders, TC_PR_ORDER)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        _insert_ordered(tr_pr, OxmlElement("w:cantSplit"), TR_PR_ORDER)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        _insert_ordered(tr_pr, OxmlElement("w:tblHeader"), TR_PR_ORDER)


def format_cell(cell, text: str, *, bold: bool = False, color: str = BODY, size: float = 10.3):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.clear()
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_table(
    doc: Document,
    headers: list[str] | None,
    rows: list[list[str]],
    widths: list[int],
    *,
    first_column_emphasis: bool = False,
    body_font_size: float = 10.3,
):
    row_count = len(rows) + (1 if headers else 0)
    table = doc.add_table(rows=row_count, cols=len(widths))
    table.style = "Table Grid"
    cursor = 0
    if headers:
        header_row = table.rows[0]
        repeat_header(header_row)
        prevent_row_split(header_row)
        for index, text in enumerate(headers):
            set_cell_shading(header_row.cells[index], TABLE_FILL)
            set_cell_borders(header_row.cells[index])
            format_cell(header_row.cells[index], text, bold=True, color=NAVY, size=10.2)
        cursor = 1
    for row_index, values in enumerate(rows, start=cursor):
        row = table.rows[row_index]
        prevent_row_split(row)
        for col_index, text in enumerate(values):
            cell = row.cells[col_index]
            set_cell_borders(cell)
            if first_column_emphasis and col_index == 0:
                set_cell_shading(cell, LIGHT_FILL)
            format_cell(
                cell,
                text,
                bold=first_column_emphasis and col_index == 0,
                color=NAVY if first_column_emphasis and col_index == 0 else BODY,
                size=body_font_size,
            )
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=PAGE_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=CELL_MARGINS_DXA,
    )
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc: Document, text: str, *, label: str | None = None):
    p = doc.add_paragraph()
    p.style = doc.styles["Callout"]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.left_indent = Pt(9)
    p.paragraph_format.right_indent = Pt(9)

    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_FILL)
    shd.set(qn("w:val"), "clear")
    _insert_ordered(p_pr, shd, P_PR_ORDER)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "6")
        border.set(qn("w:color"), BORDER)
        borders.append(border)
    _insert_ordered(p_pr, borders, P_PR_ORDER)

    if label:
        lead = p.add_run(f"{label}  ")
        set_run_font(lead, size=10.5, color=BLUE, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=11.5, color=NAVY, bold=True)
    return p


def add_body(doc: Document, text: str, *, bold_lead: str | None = None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=11, color=BODY, bold=True)
        rest = p.add_run(text[len(bold_lead) :])
        set_run_font(rest, size=11, color=BODY)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=BODY)
    return p


def add_heading(doc: Document, text: str, level: int, *, page_break_before: bool = False):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.page_break_before = page_break_before
    for run in p.runs:
        set_run_font(
            run,
            size={1: 16, 2: 13, 3: 12}[level],
            color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
            bold=True,
        )
    return p


def build_document():
    doc = Document()
    configure_document_defaults(doc)
    section = doc.sections[0]
    configure_section(section)
    configure_header_footer(section)
    configure_styles(doc)
    bullet_num_id = create_numbering(doc, decimal=False)
    number_num_id = create_numbering(doc, decimal=True)

    props = doc.core_properties
    props.title = "企业一体化经营管理平台 - 产品介绍与功能大纲"
    props.subject = "面向非技术读者的产品介绍（总体设计阶段，功能以正式发布版本为准）"
    props.author = "企业一体化经营管理平台"
    props.keywords = "企业管理, CRM, 合同, 订单, 采购, 财务, 售后, 私有化, 自动化, 能力包, MCP"
    props.comments = ""
    props.last_modified_by = "企业一体化经营管理平台"
    props.created = DOCUMENT_TIMESTAMP
    props.modified = DOCUMENT_TIMESTAMP

    # First-page customer-pack header pattern, without a decorative rule.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(2)
    run = kicker.add_run("产品介绍 / 功能大纲")
    set_run_font(run, size=10.5, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("企业一体化经营管理平台")
    set_run_font(run, size=29, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.style = doc.styles["Intro Lead"]
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("把客户、合同、订单、采购、库存、财务和售后连成一条线")
    set_run_font(run, size=13.5, color=MUTED)

    add_callout(
        doc,
        "同一份信息只录入一次，后续部门自动接力；每件事都有负责人、进度、审批和完整记录。",
        label="一句话说明",
    )

    add_body(
        doc,
        "文档状态：F-57 现行产品介绍。当前只完成总体设计、需求追踪和实施计划，F-57 功能尚未实现；实际交付必须以正式发布版本、合同和验收清单为准，详见第 9 节。",
    )

    add_heading(doc, "这是一套什么软件？", 1)
    add_body(
        doc,
        "这是一套由企业自己掌控的综合经营管理软件。它把原来分散在不同软件、Excel、邮件、聊天和纸质单据里的工作，放进同一个安全、可追溯的工作体系。",
    )
    add_body(
        doc,
        "一笔业务从商机和报价开始，一直跟踪到合同、订单、采购、收货、交付、开票、收付款和售后，再持续回到复购、续签、维保和经营改进。系统用耐久自动化把责任、效果、证据、异常与重新开启连成闭环；MCP 是受治理的工具层，服务器控制中心是权威管理面。本地模型延后开发，当前只冻结可替换 AI provider、权限、审计和隔离契约。",
    )

    add_table(
        doc,
        ["项目", "说明"],
        [
            ["使用设备", "员工 Workbench 支持 Windows、macOS、iOS 和 Android 并自适应屏幕；权威服务器控制中心只运行在 Windows Server 2022"],
            ["部署方式", "当前首版只在企业自控的 ThinkStation P340 物理机上运行一台 Windows Server 2022 作为唯一写权威；服务器外备份目标、离线轮换介质和洁净恢复能力独立配置。核心交易数据库为自管 PostgreSQL 16。未来 IaaS 只能通过新的独立认证档启用，详见第 6 节"],
            ["数据原则", "每个客户独立部署；同一客户数据库按法人行级权限和独立密钥域隔离；权威节点承载内容或可关联客户的持久数据与衍生数据全部落加密 HDD，Workbench 只允许最小、加密、可撤销、非权威缓存"],
            ["核心特点", "全链路闭环、长周期耐久自动化、动态权限、签名配置代、能力包热插拔、受控 MCP、高度定制和完整审计"],
            ["适用范围", "首版面向中国大陆业务，仅支持简体中文与人民币；不含多币种、外汇和进出口"],
        ],
        [2700, 6660],
        first_column_emphasis=True,
    )

    add_heading(doc, "适合哪些企业？", 2)
    add_body(
        doc,
        "尤其适合以合同和订单驱动经营，销售、采购、仓库、财务和售后需要紧密协同的企业，也适合多法人、多部门、多地点、重视数据保密或需要大量定制流程的组织。",
    )

    add_heading(doc, "1. 一笔订单怎样走完全程", 1)
    add_body(doc, "用最常见的一笔订单来理解这套系统：")
    flow_steps = [
        ("销售建单：", "销售录入合同，选择客户、产品、数量、价格和交期，系统自动带出已有资料，并校验价格权限、库存可用量、交期和客户信用额度。"),
        ("合同审批：", "关键条款、折扣、付款计划和附件按当前风险策略进入审批或会签，不能越权跳过；授权依据、意见、版本和附件全程留痕。"),
        ("合同生效：", "合同生效后自动生成销售订单、采购需求、项目任务、收款计划和交付节点，派生单据与合同双向可追溯。"),
        ("采购订货：", "采购查看合同来源和已批准的采购需求，按采购需求下达采购订单并可分批订货，供应商可通过门户确认订单与交期。"),
        ("收货入账：", "收货按物料、批次和序列号写入库存台账，并按采购订单不含税单价暂估入库金额；采购发票登记后回冲暂估，按发票不含税单价调整入库成本并形成应付明细。"),
        ("发票申请与开具：", "销售按合同和订单提交发票申请，经当前策略解析出的授权审批人批准后，由财务登记开具结果，回写状态与剩余可开比例并形成应收。"),
        ("交付确认：", "发货或合同交付节点确认后确认收入；仅非直运 INVENTORY 行按当前移动加权成本同步结转销货成本。DROP_SHIP 和 DIRECT_EXPENSE 行不产生本方库存或虚构销货成本腿，实际成本来自已确认且未冲销的直接采购、外购服务发票或其他权威成本事实。"),
        ("到款与付款登记：", "财务按订单和发票登记到款，按采购发票登记付款，支持分次收付款，未核销部分进入账龄；暂无可核销发票的款项先挂预收或预付，后续开票或采购发票登记时自动核销。"),
        ("售后工单：", "售后技术支持记录形成工单，关联原订单、合同、产品、批次、设备和保修，并进入客户档案。"),
        ("管理层看数：", "管理层随时查看收入、成本、交付和利润，并可按期间、客户、产品和合同下钻。"),
    ]
    for lead, rest in flow_steps:
        add_numbered(doc, lead + rest, number_num_id, bold_lead=lead, compact=True)

    add_callout(
        doc,
        "合同 -> 订单 -> 采购收货 -> 交付确认 -> 开票 -> 收付款 -> 售后 -> 经营看数，全程共用同一条业务记录。",
        label="核心闭环",
    )

    add_heading(doc, "2. 它解决的日常问题", 1)
    add_table(
        doc,
        ["常见情况", "使用平台后", "直接结果"],
        [
            ["客户资料在个人手里", "统一客户档案，历史合同、回款、投诉、设备和服务记录集中可查", "人员变化也不会丢客户"],
            ["合同、回款和附件分散", "关键条款、收付款信息、版本和附件放在一起", "履约、到期和欠款随时可查"],
            ["订单靠 Excel 或聊天转发", "合同生效后自动派生订单、采购需求、收款计划和交付节点", "少重复录入，少漏单错单"],
            ["采购不知道销售何时要货", "采购建议按合同、订单和库存不足自动形成", "减少缺料、积压和紧急采购"],
            ["赊销额度靠人工把关", "下单时自动校验客户信用额度，超额按配置阻断或转审批", "减少超额赊销带来的呆账"],
            ["发票和收款对不上", "发票申请、开具登记、收款计划与到款核销逐笔勾稽", "应收余额和账龄一目了然"],
            ["售后找不到单据出处", "工单关联原订单、合同、产品、批次、设备和保修", "处理过程清楚，责任明确"],
            ["经营报表月底手工拼", "收入、成本、交付和利润由业务单据持续汇总；常用报表的性能指标以合同约定的基线环境与验收清单为准", "管理层更早发现问题"],
        ],
        [2500, 3590, 3270],
        body_font_size=9.9,
    )

    add_heading(doc, "3. 核心业务功能", 1)

    add_heading(doc, "客户与销售", 2)
    add_body(doc, "统一管理客户档案、联系人和客户 360 视图，销售建单时自动带出客户、产品和价目资料。")
    add_bullet(doc, "销售可查看客户的历史合同、回款、投诉、设备和服务记录。", bullet_num_id)
    add_bullet(doc, "维护产品价目，录入合同和下单时校验价格权限，折扣随合同审批链审批。", bullet_num_id)
    add_bullet(doc, "支持商机、报价、跟进和转合同或订单；市场活动、销售预测和渠道佣金继续延期。", bullet_num_id)

    add_heading(doc, "合同管理", 2)
    add_body(doc, "重点把合同的三类信息管清楚：关键条款、收付款信息、合同附件。")
    add_bullet(doc, "支持模板、条款、修订版本、批注、审批、通过经认证 provider 的电子签章、实体印章、履约和义务跟踪，合同也可以合并。", bullet_num_id)
    add_bullet(doc, "支持合同续签：按原合同派生续签版本，保留与原合同的关联关系，以及原合同的履约记录、收付款计划和已派生单据的追溯链路；续签版本重新审批生效后派生新的订单、收款计划和交付节点。", bullet_num_id)
    add_bullet(doc, "支持合同到期提醒：按合同有效期、交付节点日期和收付款计划到期日生成提醒。", bullet_num_id)
    add_bullet(doc, "合同完成签署并生效后，可自动生成订单、采购需求、项目任务、收款计划和交付节点。", bullet_num_id)
    add_bullet(doc, "合同变更只影响未履行义务，由新版本和正式影响计划生成或调整下游新版本；已经交付、开票、收付款、出入库、签章或产生其他业务效果的事实保持不可变，需要时只能追加冲销、更正、退换或补偿事实。派生单据与合同双向可追溯。", bullet_num_id)
    add_bullet(doc, "合同生效属高风险操作，须重新确认身份，审批链不能越权跳过。", bullet_num_id)

    add_heading(doc, "订单管理", 2)
    add_body(doc, "销售订单只能来自已生效合同版本、已接受报价版本，或经独立审批的人工建单依据，三者必须且只能选择一项；每张订单冻结完整商业快照，并校验价格权限、库存可用量、交期和客户信用额度。")
    add_bullet(doc, "客户信用占用由应收有效未收、已交付未开票和在途订单三部分组成，同一订单金额在生命周期任一时点只落入其中一处，不能重复占用；可用额度不足时按配置阻断或转独立审批。", bullet_num_id)
    add_bullet(doc, "支持订单变更、拆分、合并、取消、分批交付、退货、换货和直运；STANDARD 与 DROP_SHIP 是首版必须完成的预发布认证目标，当前仍未实现、未认证，证据通过前不得宣称可用。寄售、订阅和租赁只保留类型化 provider seam，不提供可执行菜单、接口或营销声明。", bullet_num_id)
    add_bullet(doc, "每次变更都保留版本和审批记录，避免口头修改造成部门信息不一致。", bullet_num_id)

    add_heading(doc, "采购与供应商", 2)
    add_body(doc, "采购需求可由合同、销售订单、项目、库存补货、经审批人工需求或受控外部生产请求形成；合并、拆分和部分订购后仍保存逐来源数量并保持总量守恒。")
    add_bullet(doc, "采购按已批准采购需求下达采购订单，可合并、拆分和分批订货，覆盖询比价、授标、收货、退货、采购发票和付款申请。", bullet_num_id)
    add_bullet(doc, "记录供应商准入与资质，以及价格、交期、质量和风险信息。", bullet_num_id)
    add_bullet(doc, "供应商可通过门户确认采购订单与交期、提交送货通知、上传发票、查询收付款对账并维护自身档案。", bullet_num_id)
    add_bullet(doc, "支持询价、供应商报价、比价与选择审计；招投标和 VMI 继续延期，供应商价格、交期、质量和风险评价纳入主档。", bullet_num_id)

    add_heading(doc, "3. 核心业务功能（续）", 1)

    add_heading(doc, "库存与存货计价", 2)
    add_body(doc, "管理仓库与库存台账、收发存记录、可用量查询，以及批次与序列号标识。")
    add_bullet(doc, "收货、出库和退货按物料、批次和序列号登记，可用量直接支撑下单校验和交期判断。", bullet_num_id)
    add_bullet(doc, "存货计价采用移动加权平均一种方法：收票前按采购订单不含税单价暂估入库，收票后按发票不含税单价调整，差额对仍在库部分调整加权平均单价、对已出库部分计入当期成本；出库按加权平均单价结转金额，数量账与金额账同步更新。", bullet_num_id)
    add_bullet(doc, "退货只追加冲回事实，不重算历史：非直运 INVENTORY 销售退货始终按所关联原交付确认行的实际成本分段回收入库；物料采购退货不论是否已收票，库存一律按退货时锁后的当前移动加权账面价值出库，若退货后结存数量归零则全额出清剩余库存金额；原暂估、发票、红字和价差由 GRNI 与成本链另行勾稽。", bullet_num_id)
    add_bullet(doc, "首版不含高级 WMS 的库位、质检、销售分配库存预留、拣货、波次、调拨和盘点；售后工单的服务配件预留、领用、退回和报损属于当前范围。首版也不含先进先出、标准成本和采购费用分摊。", bullet_num_id)

    add_heading(doc, "经营财务：应收应付、开票与经营账", 2)
    add_body(doc, "业务单据按固定规则生成经营财务事实，减少重复录入；平台闭合履约、收付款、发票、成本与毛利，法定财税通过专业系统连接器承接。")
    add_bullet(doc, "应收台账按客户、合同、订单和发票记录应收明细、收款计划、到款核销和账龄；应付台账按供应商、采购订单和采购发票记录应付明细、付款申请、付款核销和账龄。", bullet_num_id)
    add_bullet(doc, "发票申请、审批、开具登记与合同收款计划的勾稽在系统内连成一条链路；开票有误时可在系统内登记作废或红字冲销。", bullet_num_id)
    add_bullet(doc, "到款和付款按订单与发票登记，支持分次收付款，一笔款项可核销多张发票或订单；银行与现金账户档案和资金流水由人工登记。", bullet_num_id)
    add_bullet(doc, "预收与预付纳入台账：收到尚无应收可核销的款项按合同收款计划挂预收账款，付出尚无应付可核销的款项挂预付账款，后续开票或采购发票登记时自动核销，核销后的余额进入应收应付台账并参与账龄。", bullet_num_id)
    add_bullet(doc, "经营账保存不可变且平衡的内部经营分录、受控经营科目映射、试算、业务子账对账和经营期间永久锁定；锁定后不反结账、不重开。迟到事实进入下一个开放经营期间，同时保留原业务日期、顺延依据和追加更正链。它不冒充法定科目、法定凭证账簿、税务申报、工资或法定年结。", bullet_num_id)
    add_bullet(doc, "经营成本按权威事实聚合：非直运库存行在交付确认时按移动加权结转销货成本；DROP_SHIP、外购服务及其他 DIRECT_EXPENSE 成本来自已确认且未冲销的采购发票或成本捕获；服务成本来自配件估价、已批准工时和费用；项目还可纳入经批准的其他经营成本。每项均保存来源、冲销和更正链，并可按合同、订单、客户和项目下钻计算毛利。", bullet_num_id)
    add_bullet(doc, "收入确认当前只有交付确认一种时点：所有适用行确认收入，但只有非直运 INVENTORY 行同步结转销货成本；DROP_SHIP 和 DIRECT_EXPENSE 行不产生虚构库存或销货成本腿。开票时形成应收和销项税额。", bullet_num_id)
    add_bullet(doc, "第一阶段按人民币经营事实和实际开票结果登记。只有权威登记表已具名的法定财税、电子发票、银行与支付能力保留受治理 provider 接口，并且必须逐个完成签名登记和认证后才能启用；工资、固定资产、预算、费用报销、合并报表、多币种和外汇仍属不支持或延期范围，不能靠通用 provider 宣称可用。", bullet_num_id)

    add_heading(doc, "售后工单与设备台账", 2)
    add_body(doc, "售后覆盖安装、维修、巡检、保养和技术支持五类工单，关联原订单、合同、产品、批次、设备、保修和服务权益。")
    add_bullet(doc, "设备台账记录设备编号或序列号、型号、所属客户、关联产品与批次、交付与安装日期和当前状态。", bullet_num_id)
    add_bullet(doc, "保修记录起止日期、保修范围和条款文本，建工单时自动读取在保状态。", bullet_num_id)
    add_bullet(doc, "客户投诉与工单进入客户 360 视图，退换修登记与订单的退货、换货打通。", bullet_num_id)
    add_bullet(doc, "工单按能力、位置、负载、SLA、回避和职责分离动态派工，记录现场照片、签字、配件、工时、成本、根因、纠正措施和回访；当前可按合同或保修规则派生服务权益并生成周期维护任务，周期账单/催收与售后知识库继续延期。", bullet_num_id)

    add_heading(doc, "报表与经营看板", 2)
    add_body(doc, "报表和看板的数据直接来自业务单据，管理层不必等月底手工汇总。")
    add_bullet(doc, "首版计划预置收入、成本、毛利、交付、应收应付账龄、采购周期、库存、服务 SLA/成本/满意度、项目风险，以及目标闭环、未知效果和自动化健康指标，并提供默认管理驾驶舱。", bullet_num_id)
    add_bullet(doc, "每个指标都保存公式版本、来源水位和证据，可按适用的期间、法人、客户、产品、合同、订单、项目、工单和责任链下钻；关闭、重开、冲销和异常不会被静默排除。", bullet_num_id)
    add_bullet(doc, "支持自定义指标、报表、看板和打印模板，并作为受审批的同一配置代发布、验证和回滚；具体编辑交互以实施验收版本为准。", bullet_num_id)
    add_bullet(doc, "报表和看板的结果继承法人、记录和字段级权限，无权查看的数据不会出现在结果中。", bullet_num_id)
    add_bullet(doc, "指标必须解释公式并下钻到来源证据；内嵌电子表格和外部 BI 语义层继续延期，定时任务通过受治理自动化执行。", bullet_num_id)

    add_heading(doc, "4. 首版模块安装范围与后续版本", 1)
    add_body(
        doc,
        "首版计划按模块交付。首版范围内的模块完成实现与验收后，可以按许可证安装、启用、停用、再启用和升级；停用只关闭该模块的界面入口、写入接口、定时任务和对外事件，历史数据继续保留，授权范围内仍可查询和审计检索。",
    )

    add_heading(doc, "首版能力概览（易读分组，非模块注册表）", 2)
    add_body(doc, "下面九组只为非技术读者理解，不是安装、许可或数据库模块登记。唯一机器可执行的内置模块闭集仍是 15 个 ModuleCode：mdm、crm、cpq、clm、sales、procure、inventory、costing、project、service、finance、ledger、invoice、portal、reporting；审批、自动化、权限、附件、搜索、MCP、provider、AI 契约和服务器控制中心属于平台公共能力，不构成第 16 个模块。")
    first_release_modules = [
        ("核心业务：", "主数据、客户与客户 360、商机与报价、合同、STANDARD/DROP_SHIP 销售、询比价采购、基础库存、项目与交付、投诉、售后工单、设备和周期维保。"),
        ("经营财务：", "应收应付、预收预付、发票、收付款、退款返款、核销、账龄、现金流、内部平衡经营分录、试算、子账对账、经营期间和毛利。"),
        ("耐久自动化与动态权限：", "以目标、责任、效果、证据、异常、闭环和周期驱动长链工作；任务按能力动态找人，角色与岗位只作模板。"),
        ("报表与平台定制：", "可定义关系型业务对象、字段、关系、表单、列表、菜单、流程、权限、报表、看板、打印模板和品牌，并作为同一签名配置代发布、回滚。"),
        ("客户与供应商门户：", "客户按白名单查看业务并确认交付、提交投诉或服务请求；供应商只使用订单交期、ASN、发票、对账和自有资料五类能力。"),
        ("受治理 MCP 与 provider：", "本地文件、Office 格式、REST/Webhook/MCP、SMTP 和 AD/LDAP 是必须完成认证的核心 provider 目标；证据通过前相应能力保持关闭。其他厂商连接器逐项取证后以签名能力包启用。MCP 不能直连数据库。"),
        ("AI 契约：", "交付可替换 AI provider、模型与工具版本、最小外发、权限、审计和隔离契约；本地模型本身延期，确定性主链不依赖 AI。"),
        ("服务器控制中心：", "权威节点提供配置代、动态权限、能力包、自动化、审计、安全、磁盘、备份、恢复和降级状态管理；它不是第五个办公客户端，也没有绕权超级管理员。"),
        ("公共基础：", "审批会签、SLA、通知、附件、全局搜索、模板、Excel 导入导出、审计和可解释错误。"),
    ]
    for lead, rest in first_release_modules:
        add_bullet(doc, lead + rest, bullet_num_id, bold_lead=lead, compact=True)

    add_heading(doc, "属于后续版本的能力", 2)
    add_body(
        doc,
        "以下十二项逐一对应权威登记的 12 个产品边界别名，不是可随意扩写的愿望清单。后续恢复必须引用对应登记项、重新裁定范围并取得证据，不能在首版通过低代码配置、插件或连接器变相实现。",
    )
    deferred_groups = [
        ("线索与市场：", "市场活动、营销漏斗、渠道佣金和销售预测延期；CRM 当前只接收类型化客户与商机输入。"),
        ("复杂报价：", "复杂产品配置器、复杂成本模型、返利和报价部分接受延期；基础报价版本属于当前范围。"),
        ("复杂采购：", "正式招投标、VMI 和复杂供应商绩效模型延期；RFQ 与询比价属于当前范围。"),
        ("特殊销售模式：", "寄售、订阅和租赁的完整销售闭环延期，只保留受治理 provider 接口且默认关闭。"),
        ("制造：", "完整 MRP、MES 和 APS 延期；外部生产系统只能通过受治理接口形成标准采购需求。"),
        ("高级仓储：", "高级 WMS 的波次、拣货、盘点、质检、销售分配库存预留、调拨和自动立库延期；基础库存及服务配件预留属于当前范围。"),
        ("深度服务经营：", "周期计费引擎、预测维护和完整 EAM 延期；服务权益、成本、一次性收费提案和周期维保任务属于当前范围。"),
        ("深度项目管理：", "完整 WBS、资源、预算变更和 EVM 延期；基础项目、风险、成本和收款节点属于当前范围。"),
        ("法定财务：", "法定总账、税务、工资和法定年结延期，由专业系统通过受治理接口承接；内部经营分录属于当前范围。"),
        ("企业专业套件：", "HR、GRC、法务、商旅、ECM、GIS、PLM、PIM 和 QMS 延期，不创建首版模块、菜单或接口。"),
        ("其他门户：", "经销商门户和员工门户延期；客户门户与供应商门户按当前精确白名单交付。"),
        ("本地智能能力：", "本地模型、OCR、RAG 和知识图谱实现延期；当前只交付 AI/provider 治理契约、受控 MCP 和权限过滤全文检索。"),
    ]
    for lead, rest in deferred_groups:
        add_bullet(doc, lead + rest, bullet_num_id, bold_lead=lead, compact=True)

    add_callout(
        doc,
        "第一阶段交付以合同履约与回款为商业切口，但技术内核是可组合、可热插拔、可恢复的治理自动化平台。本地模型不在当前交付中；受控 MCP、AI/provider 契约和服务器控制中心在当前范围内。",
        label="范围提醒",
    )

    add_heading(doc, "高度定制，但不破坏升级", 1)
    customization_items = [
        ("数据可定制：", "新增业务对象、字段、关系、编号、校验、视图和搜索。"),
        ("界面可定制：", "调整表单、列表、首页、菜单和看板（同一配置在移动端按重排规则呈现）。"),
        ("流程可定制：", "设置审批、会签、时限、提醒、升级、自动任务和跨部门动作。"),
        ("权限可定制：", "按主体、能力、法人、对象、记录、字段、条件、期限、设备、金额、状态和委托动态控制；岗位和角色只作模板。"),
        ("报表可定制：", "建立企业自己的指标、报表、打印模板和管理驾驶舱。"),
        ("品牌可定制：", "每个客户可使用自己的名称、图标、颜色和客户端包；面向外部人员的应用商店分发使用客户自有开发者账号与证书。公共商店审核连续两轮失败或超过合同约定 14 个自然日，并经客户批准后，才可切换到 Web/PWA 或其他约定形态并记录能力差异。"),
        ("扩展可安装：", "通过签名模块或受控插件，在首版已交付的能力范围内增加特殊能力；配置先验证审批，出现问题可回退。插件默认没有网络、文件、密钥和业务数据权限，需要哪些能力必须先声明，经审批后按最小权限授予。扩展可运行于 WASM、受 Job Object 隔离的签名 Windows worker，或已通过 Hyper-V/容器/资源安全证据的受控 Windows 容器；当前 P340 32GB 默认不激活容器。"),
    ]
    for lead, rest in customization_items:
        add_bullet(doc, lead + rest, bullet_num_id, bold_lead=lead, compact=True)

    add_heading(doc, "5. Excel、文档与外部连接", 1)

    add_heading(doc, "Excel 与批量数据", 2)
    add_bullet(doc, "首版计划提供 Excel 导入与导出，用于批量建档、批量录入和把查询结果带出系统。四端的权限、校验、审计和服务器结果一致；电脑端只负责建立映射、预览并启动任务，文件解析、权威校验与执行只在服务器完成。移动端只能查看、审批、暂停或恢复服务端任务，单次批量规模按设备策略受限。", bullet_num_id)
    add_bullet(doc, "为了安全，不执行 VBA、任意宏或第三方 Excel 加载项。", bullet_num_id)
    add_bullet(doc, "首版不提供 Excel 或 WPS 加载项，也不提供在表格中实时查询与提交、复杂公式、数据透视和条件格式。", bullet_num_id)

    add_heading(doc, "合同、附件与档案", 2)
    add_bullet(doc, "合同与单据附件支持模板套用、批注、版本与版本比较；DOCX 提供下载与查看。", bullet_num_id)
    add_bullet(doc, "PDF 支持查看、下载、批注、发起签署请求和归档校验；具体签章 provider 通过认证并启用后，才可执行电子签章。合同、发票、内部经营分录证据和审计证据采用只追加记录、摘要校验、签名审计和服务器外备份共同保护。首版不宣称使用经认证的 WORM（不可重写）存储，也不把内部经营记录称为法定会计凭证。", bullet_num_id)
    add_bullet(doc, "大文件支持分片上传和断点续传；上传完成只表示文件已完整进入隔离区，仍须通过类型识别、恶意内容检查和权限校验并达到 PUBLISHED 状态后，才可成为可用附件。", bullet_num_id)
    add_bullet(doc, "移动端不提供文档模板编辑、签章坐标设计或复杂版本合并；可查看 exact digest 和版本比较、提交批注、发起签署请求并执行获准审批。签章私钥与实际签署、写入效果始终在服务器或经认证 provider。", bullet_num_id)
    add_bullet(doc, "首版不提供文档在线编辑、修订与多人协作，也不提供 OCR、PDF/OFD 遮盖与格式转换、CAD 预览。", bullet_num_id)

    add_heading(doc, "接口与外部连接", 2)
    add_body(
        doc,
        "第一阶段必须完成认证的核心 provider 目标是本地文件、Excel/CSV/Word/PDF、REST/Webhook/MCP、SMTP 和 AD/LDAP；当前均不得在缺少一致性与安全证据时宣称可用。电子签章、企业微信、钉钉、飞书、Microsoft 365、WPS、银行、税务、OIDC/SAML 等使用同一签名 provider 契约，只有具体厂商通过验收后才能启用；不是全部预装的现成连接器。",
    )
    add_bullet(doc, "银行与税务侧首版不做系统对接：到款、付款与发票开具都在系统内登记，可人工录入或批量导入，闭环不依赖银企直连与税务平台。", bullet_num_id)
    add_bullet(doc, "平台提供查询类 REST 接口与 OpenAPI 说明、Webhook 与业务事件外发，以及 CSV、Excel 文件导入导出。通用 XML/SOAP/XSD 不属于首版核心格式；只有声明具体格式与 schema、通过认证的签名 provider codec，才能把特定 XML 转为导入提案或类型化命令。需要双向读写时，只能使用已签名、已审批的 MCP/provider 清单和短期授权；系统每次重新检查人员、设备、法人、权限和字段范围。", bullet_num_id)
    add_bullet(doc, "MCP 连接默认没有网络、文件、密钥或业务数据权限；每项能力须明确声明域名、端口、对象、字段、凭据引用、文件、资源、风险和审批。写工具只能调用类型化业务命令；通用 SQL、Shell、任意文件和任意网络代理永远不可开放。", bullet_num_id)
    add_bullet(doc, "对外调用按风险和幂等契约使用超时、限次退避、熔断与死信；只有已证明未执行或具备强幂等、可安全重放的调用才能自动重试。外部效果未知时必须进入 RECONCILING 或 INCIDENT，先查询 provider；高风险效果只能由双人签名处置，禁止盲重试，业务责任持续保留直至对账或补偿闭合。", bullet_num_id)

    add_heading(doc, "6. 四端使用、私有部署与安全", 1)

    add_heading(doc, "Windows、macOS、iOS、Android", 2)
    add_body(
        doc,
        "四个平台使用同一套数据、权限和业务规则：同一操作在任一端发起、审批或查询，产生的记录、校验、审计和结果相同；界面按屏幕和使用场景重排，不要求四端外观一致。电脑端适合批量操作和复杂报表，手机适合审批、查询、扫码、拍照和现场记录；权威配置只在服务器控制中心完成。",
    )
    mobile_scope_items = [
        ("移动端完整使用：", "库存台账与收发扫码、售后工单与设备台账、审批待办与站内通知。"),
        ("移动端简化使用：", "客户档案与客户 360 查询、合同条款与电子签章、销售订单与履约、采购与供应商协同、项目任务与交付节点、主数据维护与审批、全文检索；业务对象、权限和流程结果不变，交互形态与单次批量规模受限。"),
        ("移动端受设备策略限制：", "付款、退款、经营期间、敏感导出等高风险动作依当前设备、金额、状态、重新认证和职责分离策略决定；服务器配置与能力包发布不属于 Workbench。"),
        ("移动端不承载：", "移动端不加载能力包、WASM、原生插件或动态下载的可执行扩展代码；只下发签名 UI schema、规则数据、模板和静态资源，相机、扫码、拍照、触控签字等设备能力随已签名应用版本发布。"),
    ]
    for lead, rest in mobile_scope_items:
        add_bullet(doc, lead + rest, bullet_num_id, bold_lead=lead, compact=True)
    add_bullet(doc, "所有业务写入经权威端校验后生效。最高安全档默认不启用业务投影离线读取；只有签名设备策略逐对象、逐字段显式开放时，才可在最长 24 小时的有界期限内读取最小、加密、可撤销的非权威投影。可保存表单与附件草稿、现场证据和待提交意图；恢复连接后服务器重验权限、配置代、记录版本和幂等。付款、最终审批、合同生效、库存权威、权限和配置不能离线生效。", bullet_num_id)
    add_bullet(doc, "合同生效、付款与退款、开票与红冲、经营分录更正、经营期间锁定、迟到事实顺延例外和敏感数据导出属于高风险操作，必须按现行风险策略重新认证并进入相应审批，审批人不得与发起人为同一人，四端口径一致。", bullet_num_id)

    add_heading(doc, "企业自己选择部署位置", 2)
    add_bullet(doc, "首版只有一台运行在企业自控 ThinkStation P340 物理机上的 Windows Server 2022 单写权威，不依赖第二台应用服务器。最高安全生产仍必须另配服务器外连续备份目标、至少两块轮换离线加密介质和洁净 Windows 恢复能力。未来可增加客户自控境内 IaaS 独立认证档，但当前不实现、不接受；必须先发布新的架构与认证版本，并分别证明驻留、vTPM、底层 HDD、缓存、快照、运维副本、备份域和安全关机边界，不能复用 P340 证据。", bullet_num_id)
    add_bullet(doc, "首版核心交易数据库只支持并只认证 PostgreSQL 16，由客户在同一台服务器上自主管理；不使用云托管数据库。", bullet_num_id)
    add_bullet(doc, "企业已有的 Oracle、SQL Server、MySQL 可作为外部数据源接入，不作为核心交易数据库。", bullet_num_id)
    add_bullet(doc, "核心安装不依赖云托管数据库、消息、遥测、更新或厂商控制面；本地、自建、私有云或已有服务都必须实现同一受控 provider 契约，逐项认证后才可连接。", bullet_num_id)
    add_bullet(doc, "SSD 只放 Windows、程序、静态资源和可重建依赖；数据库、WAL、附件、索引、日志、审计、导出、临时业务文件和所有衍生数据全部落到加密 HDD。", bullet_num_id)
    add_bullet(doc, "真实客户数据和一切可关联客户的衍生数据只允许在中国大陆境内处理和持久化，覆盖数据库/WAL、附件、索引、日志、审计、导出、临时文件、备份与离线轮换、恢复材料元数据、监控、支持诊断、provider 输入输出和可关联遥测；地点未知、证据过期或存在跨境路径时，相应能力失败关闭。", bullet_num_id)

    add_heading(doc, "高保密设计", 2)
    security_items = [
        ("权限细：", "由主体、能力、数据范围、条件、期限、设备、金额、状态和委托共同决定，可控制到法人、单条记录和单个字段；策略默认拒绝，没有绕权超级管理员。"),
        ("职责分开：", "系统、数据、安全、审计和密钥管理员相互制约；申请人不可自审，审批链不可越权跳过。"),
        ("客户持钥：", "传输、数据库、附件、备份和设备缓存均受加密保护，每个法人使用独立密钥域；客户可选择部署在自身环境、由客户控制并以 TPM 包装且配有独立恢复材料的密钥服务，也可连接客户 HSM、KMS 或已认证的既有服务。厂商默认不能取得主密钥或解密生产数据。"),
        ("控制外泄：", "敏感字段由服务器按批准清单裁剪，动态水印、导出审批和审计在三类端口强制执行。只有受支持、受管且合规的 Windows/macOS 与 iOS/Android 原生端，才按 OS/MDM 能力强制剪贴板、分享、打印、截图/录屏和受管文件失效；不合规或能力不足时降级为只读，并禁止高密级、离线或下载。浏览器门户只承诺脱敏、水印、导出审批和审计；打印、剪贴板与已下载文件失效仅尽力限制，也不宣称能够阻止外部相机。"),
        ("完整审计：", "关键查看、修改、审批、导出和系统调用留下防篡改证据，审计记录只追加、可逐条验证。"),
        ("可恢复：", "第一阶段是单写权威，不宣称高可用。上线必须同时具备独立故障域内 HDD-only 的服务器外追加式自动增量目标（关闭或证明为空的 SSD 客户数据缓存）、至少两块交替且平时断开的加密离线轮换 HDD、UPS、分域恢复材料，以及在洁净 Windows 主机的容量达标加密 HDD 工作区完成的完整恢复演练。"),
    ]
    for lead, rest in security_items:
        add_bullet(doc, lead + rest, bullet_num_id, bold_lead=lead)

    add_callout(
        doc,
        "云服务器仅是未来认证选项，当前首版只接受 P340；它不等于共享 SaaS，每个客户仍有独立系统、数据、密钥和升级通道。",
        label="重要区别",
    )

    add_heading(doc, "勒索软件与恢复边界", 2)
    add_bullet(doc, "服务器外、独立故障域、HDD-only、加密且追加式的自动增量目标，与至少两块交替使用、平时物理断开的加密离线轮换 HDD 必须同时存在；连续目标的 SSD 缓存必须关闭或经取证始终不含客户字节，洁净恢复主机也必须使用容量足以容纳实际恢复集、校验空间和增长余量的加密 HDD 工作区。两层不是二选一，同服务器目录、同盘副本或仅有云盘快照不能冒充灾难副本。", bullet_num_id)
    add_bullet(doc, "日常备份身份只能创建新副本和完成必要校验，不能删除、覆盖、改名、修改权限或改变保留策略；恢复身份和到期处置身份另行保管，关键处置由两人批准。", bullet_num_id)
    add_bullet(doc, "保护条件缺失、检查结果不确定或完整恢复演练失败时，系统会保留不可忽略的风险提示并阻止发布。", bullet_num_id)
    add_bullet(doc, "当前 P340 的单块 1TB HDD 只能是 SINGLE_DISK_DEGRADED_PRODUCTION 候选：UPS、两层备份、洁净恢复、20 人实机混合负载和同一候选版本连续 72 小时稳定运行证据未全部通过前不得录入真实客户数据；任何版本、硬件或关键配置变化都会使该稳定性证据失效并需重测。后续优先升级两块匹配企业 HDD 的经验证 RAID1 和 64GB 内存。", bullet_num_id)

    add_heading(doc, "7. 常见人物模板能得到什么", 1)
    add_body(doc, "下表只是方便理解的默认人物与工作台模板，不是固定岗位或权限边界。实际权限和任务分配始终由当前动态策略解析。")
    add_table(
        doc,
        ["人物模板", "不再困扰于", "直接得到"],
        [
            ["销售", "客户资料分散、交期不清、回款难追", "客户全景、合同与订单进度、开票与到款状态、信用额度校验"],
            ["采购", "合同要买什么不清楚、付款进度靠问", "按合同下达的采购订单与分批订货、供应商档案、付款申请与进度"],
            ["仓库", "账实不符、批次不清、出入库靠纸", "实时库存台账、扫码收发、批次与序列号追溯"],
            ["财务", "重复录入、收付款对不上、内部经营记录来源不清", "应收应付台账、分次收付款核销、可追溯的内部经营分录、试算和子账对账"],
            ["售后", "找不到历史、工单没人跟", "关联订单、合同、产品、批次、设备和保修的工单，以及客户服务历史"],
            ["管理层", "报表滞后、口径不一、无法看到全局", "合同与发票申请的审批入口，收入、成本、交付、利润的经营看板"],
        ],
        [1500, 3770, 4090],
    )

    add_heading(doc, "8. 最关键的产品定位", 1)
    add_body(
        doc,
        "首版要解决的是一件具体的事：一笔业务从销售建单和合同审批开始，经采购订货、收货入库、发票申请与开具登记、到款与付款登记、交付确认和售后工单，一直走到管理层看到收入、成本、交付和利润。当前范围内的经营事实、责任、异常和证据在同一套系统闭环，不再依赖线下影子台账；法定财税等明确延期能力仍通过受控专业系统连接器承接。",
    )
    add_body(
        doc,
        "第一阶段不追求替换所有专业系统。上文十二项已登记的后续边界继续保持关闭；客户门户、供应商门户、受控 MCP、AI/provider 契约和服务器控制中心纳入当前设计范围。企业差异通过签名配置代、行业能力包和受控插件实现，不为每个客户复制不可升级的内核分支。",
    )
    add_body(
        doc,
        "核心后台采用 Rust 构建，强调稳定、安全和长期可维护；每个客户拥有独立部署、独立数据、独立密钥和独立升级通道。",
    )

    add_callout(
        doc,
        "首版目标：让合同、订单、货物、资金、发票和服务记录在同一套可信记录中衔接，管理层随时看到同一口径的经营结果。",
        label="产品价值",
    )

    add_heading(doc, "9. 版本与适用范围说明", 1)
    add_body(
        doc,
        "本节说明本文档的效力范围。产品处于分阶段研发过程中，文中所列功能不代表当前均已可用，实际交付内容以正式发布版本、合同和验收清单为准。",
    )
    scope_items = [
        ("阶段与版本：", "产品处于总体设计与分阶段研发阶段，不设置固定的发布日期；首版范围内的全部模块完成并通过验收后才对外发布首版。"),
        ("交付依据：", "每次交付的模块范围、版本和验收标准以双方签署的合同与验收清单为准，本文档不构成功能交付承诺。"),
        ("第一阶段范围：", "CRM 商机报价、合同、STANDARD/DROP_SHIP 销售、询比价采购、基础库存、项目交付、投诉售后、经营财务、客户与供应商门户、报表、耐久自动化、动态权限、平台定制、受控 MCP、AI/provider 契约、服务器控制中心及四端 Workbench。"),
        ("后续版本：", "上文十二项已登记的产品边界，以及边界登记 DEF-006 的主主/双活/多写和 DEF-007 的 PostgreSQL 外权威数据库，均不在当前交付中；没有登记与证据不得新增宣称。"),
        ("语言与币种：", "首版固定 zh-CN、CNY 和 Asia/Shanghai 业务显示；权威时间戳保存为 UTC，持续时间、租约与超时使用 monotonic clock。不支持多币种、外汇、进出口、报关、信用证和产品内容多语言。"),
        ("联网要求：", "最高安全档默认不启用业务投影离线读取；经签名设备策略逐对象、逐字段显式开放时，只允许有界、加密、可撤销的最小非权威投影。草稿、附件、现场证据和待提交意图可暂存，所有权威写入由服务器重验后生效，高风险动作不得离线生效。"),
        ("人工确认：", "合同生效、付款与退款、开票与红冲、经营分录更正、经营期间锁定、迟到事实顺延例外和敏感数据导出属于高风险操作，必须按现行风险策略重新认证并进入相应审批；审批链不可越权跳过，审批人不得与发起人为同一人。"),
        ("数据库与部署：", "首版唯一权威数据库认证目标是在 ThinkStation P340 上同机自管 PostgreSQL 16，并以一台 Windows Server 2022 作为单写权威，同时独立配置服务器外备份与洁净恢复能力。取得实际部署证书前不得称为已认证；权威节点全部客户及衍生持久数据必须落加密 HDD，核心运行不依赖 Linux、WSL、Kubernetes、共享 SaaS 或厂商云控制面。IaaS 仅是未来独立 profile 扩展缝，当前选择必须失败关闭。"),
        ("外部连接：", "必须完成认证的核心 provider 目标为本地文件、Office 格式、REST/Webhook/MCP、SMTP 和 AD/LDAP；证据通过前保持关闭。其他身份、消息、签章、银行、税务和办公厂商按签名 provider 契约逐项取证后启用。"),
        ("性能与容量：", "ThinkStation P340 i5-10500、32GB、256GB SSD、单 1TB HDD 以约 20 名活跃用户为实机认证目标，不是登录硬上限或已达成承诺；恢复时间按数据量和实测证书签发。"),
    ]
    for lead, rest in scope_items:
        add_bullet(
            doc,
            lead + rest,
            bullet_num_id,
            bold_lead=lead,
            compact=True,
            font_size=9.5,
            line_spacing=1.05,
        )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    sanitize_docx_package(
        OUTPUT,
        package_timestamp=(
            DOCUMENT_TIMESTAMP.year,
            DOCUMENT_TIMESTAMP.month,
            DOCUMENT_TIMESTAMP.day,
            DOCUMENT_TIMESTAMP.hour,
            DOCUMENT_TIMESTAMP.minute,
            DOCUMENT_TIMESTAMP.second,
        ),
    )
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
